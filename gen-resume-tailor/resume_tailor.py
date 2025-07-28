from langchain_groq import ChatGroq
from langchain.document_loaders import WebBaseLoader
# ... other imports from your original file, EXCEPT streamlit ...
from docx import Document
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
import os
import json
from typing import Dict, List, Optional
from dotenv import load_dotenv
import numpy as np
import pandas as pd

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

class ResumeTailor:
    def __init__(self):
        """Initialize the ResumeTailor with necessary components."""
        if not GROQ_API_KEY:
            raise ValueError("GROQ_API_KEY not found in environment variables")
            
        self.llm = ChatGroq(
            model="llama-3.1-8b-instant",
            groq_api_key=GROQ_API_KEY,
            temperature=0.1,
            max_tokens=None,
            timeout=None,
            max_retries=2
        )
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text content from a PDF file using PyPDF2."""
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text content from a DOCX file."""
        doc = Document(docx_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def parse_job_description(self, job_text: str, is_url: bool = True) -> Dict:
        """Extract job details from the provided URL or text."""
        if is_url:
            loader = WebBaseLoader(job_text)
            data = loader.load()
            job_content = data[0].page_content
        else:
            job_content = job_text
        
        # Use LLM to extract structured information
        prompt = f"""Please analyze the following job description and extract key information in JSON format.
        Return ONLY a JSON object with the following structure:
        {{
            "skills": ["skill1", "skill2", ...],
            "experience": "experience requirements",
            "responsibilities": ["responsibility1", "responsibility2", ...],
            "company": "company name",
            "title": "job title"
        }}

        Job Description:
        {job_content}
        """
        
        response = self.llm.invoke(prompt)
        try:
            # Extract content and clean it to ensure it's valid JSON
            content = str(response.content).strip()
            if content.startswith("```json"):
                content = content.split("```json")[1]
            if content.endswith("```"):
                content = content.rsplit("```", 1)[0]
            return json.loads(content.strip())
        except json.JSONDecodeError:
            # Fallback structure if parsing fails
            return {
                "skills": [],
                "experience": "Not specified",
                "responsibilities": [],
                "company": "Not specified",
                "title": "Not specified"
            }
    
    def match_skills(self, resume_text: str, job_requirements: Dict) -> Dict:
        """Match resume skills with job requirements using semantic similarity."""
        try:
            # Ensure skills is a list and not empty
            if not job_requirements.get('skills') or not isinstance(job_requirements['skills'], list):
                return {
                    'matched_skills': [],
                    'missing_skills': []
                }
            
            # Common tech variations dictionary
            tech_variations = {
                'ml': ['machine learning', 'ml'],
                'ai': ['artificial intelligence', 'ai'],
                'llm': ['large language model', 'llm', 'llama', 'gpt', 'language model'],
                'nlp': ['natural language processing', 'nlp'],
                'js': ['javascript', 'js'],
                'ts': ['typescript', 'ts'],
                'py': ['python', 'py'],
                'react': ['reactjs', 'react.js', 'react'],
                'node': ['nodejs', 'node.js', 'node'],
                'db': ['database', 'db'],
                'ui': ['user interface', 'ui'],
                'ux': ['user experience', 'ux'],
                'api': ['apis', 'api', 'restful', 'rest'],
                'aws': ['amazon web services', 'aws'],
                'gcp': ['google cloud platform', 'gcp'],
                'azure': ['microsoft azure', 'azure'],
                'k8s': ['kubernetes', 'k8s'],
                'ci/cd': ['continuous integration', 'continuous deployment', 'ci/cd', 'cicd'],
                'oop': ['object oriented programming', 'object-oriented', 'oop'],
                'cv': ['computer vision', 'cv']
            }
            
            # Preprocess resume text
            resume_text_lower = resume_text.lower()
            resume_sentences = [sent.strip() for sent in resume_text.split('.')]
            
            # First pass: Direct keyword matching with variations
            matched_skills = []
            remaining_skills = []
            
            for skill in job_requirements['skills']:
                skill_lower = skill.lower()
                # Check if this skill has known variations
                variations = []
                for var_key, var_list in tech_variations.items():
                    if skill_lower in var_list:
                        variations.extend(var_list)
                    # Also check if any variation is in the skill name
                    for var in var_list:
                        if var in skill_lower:
                            variations.extend(var_list)
                
                # Add common text variations
                variations.extend([
                    skill_lower,
                    skill_lower.replace(' ', ''),
                    skill_lower.replace('-', ''),
                    skill_lower.replace('.', ''),
                    skill_lower.replace('/', '')
                ])
                
                # Remove duplicates and empty strings
                variations = list(set(filter(None, variations)))
                
                if any(var in resume_text_lower for var in variations):
                    matched_skills.append(skill)
                else:
                    remaining_skills.append(skill)
            
            # Second pass: Semantic matching for remaining skills
            if remaining_skills:
                # Prepare embeddings for remaining skills and their variations
                skill_texts = []
                skill_map = {}  # Map expanded texts back to original skills
                
                for skill in remaining_skills:
                    skill_lower = skill.lower()
                    variations = []
                    # Add known variations
                    for var_list in tech_variations.values():
                        if any(var in skill_lower for var in var_list):
                            variations.extend(var_list)
                    # Add the original skill
                    variations.append(skill_lower)
                    # Remove duplicates
                    variations = list(set(variations))
                    
                    for var in variations:
                        skill_texts.append(var)
                        skill_map[var] = skill
                
                # Convert to embeddings
                skill_embeddings = self.embedding_model.encode(skill_texts)
                resume_embeddings = self.embedding_model.encode(resume_sentences)
                
                # Calculate similarities
                similarities = resume_embeddings @ skill_embeddings.T
                max_similarities = np.max(similarities, axis=0)
                
                # Use a moderate threshold for semantic matching
                threshold = 0.6  # Higher threshold for more precise matching
                
                matched_variations = set()
                for idx, skill_text in enumerate(skill_texts):
                    if max_similarities[idx] > threshold:
                        original_skill = skill_map[skill_text]
                        if original_skill not in matched_skills:
                            matched_skills.append(original_skill)
                            matched_variations.add(skill_text)
            
            # Get missing skills
            missing_skills = [skill for skill in job_requirements['skills'] if skill not in matched_skills]
            
            return {
                'matched_skills': matched_skills,
                'missing_skills': missing_skills
            }
        except Exception as e:
            # st.error(f"Error in skill matching: {str(e)}")
            return {
                'matched_skills': [],
                'missing_skills': []
            }
    def generate_cover_letter(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> str:
        """Generate a personalized cover letter based on the resume and job requirements."""
        prompt = f"""Write a professional cover letter for a job application following this specific format and guidelines.
        
        REQUIREMENTS:
        1. STRUCTURE:
           [Your Full Name]
           [Your Address]
           [City, State ZIP]
           [Email]
           [Phone]
           
           [Date]
           
           [Hiring Manager's Name/Title]
           [Company Name]
           [Company Address]
           [City, State ZIP]
           
           Dear [Hiring Manager's Name/Title],
           
           [Body Paragraphs]
           
           Sincerely,
           [Your Name]
        
        2. CONTENT GUIDELINES:
           - Opening: Express enthusiasm for the role and company
           - Body Paragraph 1: Match your skills to job requirements
           - Body Paragraph 2: Specific achievements that demonstrate value
           - Body Paragraph 3: Company knowledge and cultural fit
           - Closing: Clear call to action
        
        CONTEXT:
        Role: {job_requirements.get('title', '[Position]')}
        Company: {job_requirements.get('company', '[Company]')}
        Required Skills: {', '.join(job_requirements.get('skills', [])[:5])}
        Experience Needed: {job_requirements.get('experience', 'Not specified')}
        Your Matched Skills: {', '.join(skill_matches['matched_skills'])}
        
        Resume Context:
        {resume_text}

        Generate a complete cover letter following the exact format above. Keep it concise and focused on key achievements and relevant skills.
        """
        
        response = self.llm.invoke(prompt)
        return str(response.content)
    
    def tailor_resume(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Generate a tailored resume using the LLM and provide improvement analysis."""
        resume_prompt = f"""You are an expert ATS optimization specialist. Rewrite the following resume to maximize its ATS score while maintaining readability.
        The goal is to significantly improve the resume's ATS score by incorporating job-specific keywords and requirements.
        Keep in mind that dont add any skills that are not explicitly mentioned in the job requirements.
        keep in mind that dont add any of those things in the resume which are not already present in the resume.

        Job Requirements to Target:
        1. Required Skills: {job_requirements.get('skills', [])}
        2. Experience Needed: {job_requirements.get('experience', 'Not specified')}
        3. Responsibilities: {job_requirements.get('responsibilities', [])}

        Current Status:
        - Matched Skills: {skill_matches['matched_skills']}
        - Missing Skills: {skill_matches['missing_skills']}

        Optimization Requirements:
        1. Keyword Integration:
           - Add ALL missing required skills with relevant context
           - Place important keywords in prominent positions
           - Use exact phrases from job requirements
           - Maintain optimal keyword density (5-8%)
        
        2. Format Optimization:
           - Use clear section headers: Summary, Experience, Skills, Education
           - Start bullets with strong action verbs
           - Ensure consistent formatting
           - Use standard bullet points
        
        3. Content Enhancement:
           - Add quantifiable metrics to achievements
           - Highlight experience matching job requirements
           - Emphasize transferable skills
           - Use industry-standard terminology
        
        4. ATS Guidelines:
           - Use full terms before abbreviations
           - Avoid tables, columns, and graphics
           - Use standard job titles
           - Place keywords near the start of bullet points

        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Return ONLY the optimized resume text. Ensure EVERY required skill and responsibility is addressed.
        """
        
        tailored_resume = str(self.llm.invoke(resume_prompt).content)
        
        # Then, get the analysis separately
        analysis_prompt = f"""Analyze how the resume matches the job requirements and provide a detailed improvement analysis.
        
        IMPORTANT RULES:
        1. ONLY mention skills that are EXPLICITLY stated in the resume
        2. DO NOT make assumptions about skills not directly mentioned
        3. DO NOT infer skills from project descriptions unless explicitly stated
        4. If a skill is missing, list it in missing skills, do not try to find similar alternatives
        5. For matched skills, quote the exact text from resume that demonstrates the skill
        
        Focus on these aspects:
        1. Skills alignment - EXACT matches only
        2. Experience relevance - DIRECT matches only
        3. Achievement emphasis - ACTUAL achievements mentioned
        4. Missing keywords - List ALL required skills not found in resume
        
        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Currently Matched Skills (verified): {skill_matches['matched_skills']}
        Currently Missing Skills (verified): {skill_matches['missing_skills']}
        
        Return a JSON object with this exact structure:
        {{
            "improvements": [
                "specific improvements needed based on ACTUAL gaps"
            ],
            "skills_analysis": {{
                "matched": [
                    "ONLY skills explicitly found in resume with exact quotes"
                ],
                "missing": [
                    "ONLY skills from job requirements that are completely absent from resume"
                ]
            }},
            "achievement_emphasis": [
                "ONLY quantifiable achievements actually present in resume"
            ],
            "keyword_optimization": [
                "ONLY keywords from job requirements that should be added"
            ]
        }}
        
        Remember: Do not infer, assume, or suggest skills that are not explicitly stated in the resume.
        """
        
        analysis_response = self.llm.invoke(analysis_prompt)
        try:
            content = str(analysis_response.content).strip()
            # More robust JSON extraction
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            content = content.strip()
            
            try:
                analysis_result = json.loads(content)
                analysis_result["tailored_resume"] = tailored_resume
                return analysis_result
            except json.JSONDecodeError as json_err:
                return {
                    "improvements": ["Error analyzing improvements"],
                    "skills_analysis": {
                        "matched": ["Error analyzing matched skills"],
                        "missing": ["Error analyzing missing skills"]
                    },
                    "achievement_emphasis": ["Error analyzing achievements"],
                    "keyword_optimization": ["Error analyzing keywords"],
                    "tailored_resume": tailored_resume
                }
        except Exception as e:
            return {
                "improvements": ["Error analyzing improvements"],
                "skills_analysis": {
                    "matched": ["Error analyzing matched skills"],
                    "missing": ["Error analyzing missing skills"]
                },
                "achievement_emphasis": ["Error analyzing achievements"],
                "keyword_optimization": ["Error analyzing keywords"],
                "tailored_resume": tailored_resume
            }
    
    def generate_docx(self, tailored_content: str) -> Document:
        """Convert the tailored content into a DOCX file."""
        doc = Document()
        # Only include the actual resume content, no analysis
        content = str(tailored_content).strip()
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        return doc
    
    def generate_cold_email(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> str:
        """Generate a personalized cold email based on the resume and job requirements."""
        prompt = f"""Write a concise, professional cold email for a job application using the following details.
        Format exactly as shown, with clear sections and no commentary.
        
        REQUIREMENTS:
        1. FORMAT:
           - Subject line: Specific role and 1-2 key qualifications
           - Greeting: Professional and personalized
           - Body: 3-4 short paragraphs
           - Closing: Professional with clear call to action
        
        2. CONTENT GUIDELINES:
           - First paragraph: Introduction and mention 2-3 most relevant matching skills
           - Second paragraph: ONE specific, quantified achievement
           - Final paragraph: Brief call to action
           - Maximum 200 words
        
        CONTEXT:
        Position: {job_requirements.get('title', '[Position]')}
        Company: {job_requirements.get('company', '[Company]')}
        Key Requirements: {', '.join(job_requirements.get('skills', [])[:5])}
        Matched Skills: {', '.join(skill_matches['matched_skills'][:5])}
        
        Resume Details:
        {resume_text}
        
        EXAMPLE FORMAT:
        Subject: [Role] Application - [Key Qualification]
        
        Dear [Name],
        
        [Email Body]
        
        Best regards,
        [Full Name]
        [Contact Info]
        """
        
        response = self.llm.invoke(prompt)
        return str(response.content)

    def calculate_ats_score(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Calculate a comprehensive ATS score for the resume based on multiple factors."""
        prompt = f"""You are an ATS (Applicant Tracking System) expert. Analyze the resume against the job requirements and calculate scores.

        Rules for scoring:
        1. All scores must be integers (whole numbers)
        2. Each section score must not exceed its maximum value
        3. Total score must be the sum of all section scores
        4. Compare directly against job requirements
        5. Higher scores for exact keyword matches from requirements

        Scoring Criteria:
        1. Keyword Match (30 points max):
           - Award points for exact matches with job requirements
           - Check keyword frequency and placement
           - Required skills present: {job_requirements.get('skills', [])}
           - Current matched skills: {skill_matches['matched_skills']}

        2. Experience Alignment (25 points max):
           - Compare against required experience: {job_requirements.get('experience', 'Not specified')}
           - Check for relevant role titles
           - Evaluate described responsibilities against: {job_requirements.get('responsibilities', [])}

        3. Skills Match (25 points max):
           - Technical skills alignment
           - Soft skills presence
           - Skills context and application

        4. Education Relevance (10 points max):
           - Required education level match
           - Field of study relevance
           - Certifications value

        5. Format & Organization (10 points max):
           - Standard section headers
           - Bullet point structure
           - Content readability

        Resume to analyze:
        {resume_text}

        Job Requirements:
        {json.dumps(job_requirements, indent=2)}

        Return a JSON object with this exact structure:
        {{
            "total_score": <integer 0-100>,
            "section_scores": {{
                "keyword_match": {{
                    "score": <integer 0-30>,
                    "max": 30,
                    "details": ["<specific keywords found>", "<specific keywords missing>"]
                }},
                "experience": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<specific experience matches>", "<experience gaps>"]
                }},
                "skills": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<matched skills details>", "<missing skills impact>"]
                }},
                "education": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<education alignment details>"]
                }},
                "format": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<format strengths>", "<format improvements needed>"]
                }}
            }},
            "improvement_suggestions": [
                "<actionable suggestion 1>",
                "<actionable suggestion 2>",
                "<actionable suggestion 3>"
            ],
            "keyword_density": {{
                "<actual keyword from job requirements>": <integer frequency>
            }}
        }}
        """
        
        try:
            response = self.llm.invoke(prompt)
            content = str(response.content).strip()
            # Extract JSON content
            content = self._extract_json_content(content)
            
            # Parse and validate
            result = self._validate_ats_score(json.loads(content))
            return result
            
        except Exception as e:
            # st.error(f"Error in ATS scoring: {str(e)}")
            return self._get_default_ats_score()
    
    def _extract_json_content(self, content: str) -> str:
        """Extract JSON content from LLM response."""
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
        return content.strip()
    
    def _validate_ats_score(self, score_data: Dict) -> Dict:
        """Validate and fix ATS score data."""
        # Validate scores are integers and within range
        score_data["total_score"] = int(score_data["total_score"])
        
        # Ensure section scores are valid
        max_scores = {
            'keyword_match': 30,
            'experience': 25,
            'skills': 25,
            'education': 10,
            'format': 10
        }
        
        for section, max_score in max_scores.items():
            if section in score_data["section_scores"]:
                data = score_data["section_scores"][section]
                data["score"] = min(int(data["score"]), max_score)
                data["max"] = max_score
        
        # Recalculate total score
        total = sum(data["score"] for data in score_data["section_scores"].values())
        score_data["total_score"] = total
        
        return score_data
    
    def _get_default_ats_score(self) -> Dict:
        """Return default ATS score structure for error cases."""
        return {
            "total_score": 0,
            "section_scores": {
                "keyword_match": {"score": 0, "max": 30, "details": ["Error analyzing keywords"]},
                "experience": {"score": 0, "max": 25, "details": ["Error analyzing experience"]},
                "skills": {"score": 0, "max": 25, "details": ["Error analyzing skills"]},
                "education": {"score": 0, "max": 10, "details": ["Error analyzing education"]},
                "format": {"score": 0, "max": 10, "details": ["Error analyzing format"]}
            },
            "improvement_suggestions": ["Unable to generate suggestions due to an error"],
            "keyword_density": {}
        }
