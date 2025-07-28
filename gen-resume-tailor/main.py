import streamlit as st
from langchain_groq import ChatGroq
from langchain.document_loaders import WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from bs4 import BeautifulSoup
import requests
from docx import Document
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
import os
import json
from typing import Dict, List, Optional
from dotenv import load_dotenv
import numpy as np
import pandas as pd
from resume_form import ResumeForm
from resume_generator import ResumeGenerator
from workflow_manager import WorkflowManager
# from email_template import EmailTemplate

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

class ResumeTailor:
    def __init__(self):
        """Initialize the ResumeTailor with necessary components."""
        if not GROQ_API_KEY:
            raise ValueError("GROQ_API_KEY not found in environment variables")
            
        self.llm = ChatGroq(
            model="qwen-2.5-32b",
            groq_api_key=GROQ_API_KEY,
            temperature=0.3,
            max_tokens=None,
            timeout=None,
            max_retries=2
        )
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text content from a PDF file using PyPDF2."""
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text content from a DOCX file."""
        doc = Document(docx_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def parse_job_description(self, job_text: str, is_url: bool = True) -> Dict:
        """Extract job details from the provided URL or text."""
        if is_url:
            loader = WebBaseLoader(job_text)
            data = loader.load()
            job_content = data[0].page_content
        else:
            job_content = job_text
        
        # Use LLM to extract structured information
        prompt = f"""Please analyze the following job description and extract key information in JSON format.
        Return ONLY a JSON object with the following structure:
        {{
            "skills": ["skill1", "skill2", ...],
            "experience": "experience requirements",
            "responsibilities": ["responsibility1", "responsibility2", ...],
            "company": "company name",
            "title": "job title"
        }}

        Job Description:
        {job_content}
        """
        
        response = self.llm.invoke(prompt)
        try:
            # Extract content and clean it to ensure it's valid JSON
            content = str(response.content).strip()
            if content.startswith("```json"):
                content = content.split("```json")[1]
            if content.endswith("```"):
                content = content.rsplit("```", 1)[0]
            return json.loads(content.strip())
        except json.JSONDecodeError:
            # Fallback structure if parsing fails
            return {
                "skills": [],
                "experience": "Not specified",
                "responsibilities": [],
                "company": "Not specified",
                "title": "Not specified"
            }
    
    def match_skills(self, resume_text: str, job_requirements: Dict) -> Dict:
        """Match resume skills with job requirements using semantic similarity."""
        try:
            # Ensure skills is a list and not empty
            if not job_requirements.get('skills') or not isinstance(job_requirements['skills'], list):
                return {
                    'matched_skills': [],
                    'missing_skills': []
                }
            
            # Common tech variations dictionary
            tech_variations = {
                'ml': ['machine learning', 'ml'],
                'ai': ['artificial intelligence', 'ai'],
                'llm': ['large language model', 'llm', 'llama', 'gpt', 'language model'],
                'nlp': ['natural language processing', 'nlp'],
                'js': ['javascript', 'js'],
                'ts': ['typescript', 'ts'],
                'py': ['python', 'py'],
                'react': ['reactjs', 'react.js', 'react'],
                'node': ['nodejs', 'node.js', 'node'],
                'db': ['database', 'db'],
                'ui': ['user interface', 'ui'],
                'ux': ['user experience', 'ux'],
                'api': ['apis', 'api', 'restful', 'rest'],
                'aws': ['amazon web services', 'aws'],
                'gcp': ['google cloud platform', 'gcp'],
                'azure': ['microsoft azure', 'azure'],
                'k8s': ['kubernetes', 'k8s'],
                'ci/cd': ['continuous integration', 'continuous deployment', 'ci/cd', 'cicd'],
                'oop': ['object oriented programming', 'object-oriented', 'oop'],
                'cv': ['computer vision', 'cv']
            }
            
            # Preprocess resume text
            resume_text_lower = resume_text.lower()
            resume_sentences = [sent.strip() for sent in resume_text.split('.')]
            
            # First pass: Direct keyword matching with variations
            matched_skills = []
            remaining_skills = []
            
            for skill in job_requirements['skills']:
                skill_lower = skill.lower()
                # Check if this skill has known variations
                variations = []
                for var_key, var_list in tech_variations.items():
                    if skill_lower in var_list:
                        variations.extend(var_list)
                    # Also check if any variation is in the skill name
                    for var in var_list:
                        if var in skill_lower:
                            variations.extend(var_list)
                
                # Add common text variations
                variations.extend([
                    skill_lower,
                    skill_lower.replace(' ', ''),
                    skill_lower.replace('-', ''),
                    skill_lower.replace('.', ''),
                    skill_lower.replace('/', '')
                ])
                
                # Remove duplicates and empty strings
                variations = list(set(filter(None, variations)))
                
                if any(var in resume_text_lower for var in variations):
                    matched_skills.append(skill)
                else:
                    remaining_skills.append(skill)
            
            # Second pass: Semantic matching for remaining skills
            if remaining_skills:
                # Prepare embeddings for remaining skills and their variations
                skill_texts = []
                skill_map = {}  # Map expanded texts back to original skills
                
                for skill in remaining_skills:
                    skill_lower = skill.lower()
                    variations = []
                    # Add known variations
                    for var_list in tech_variations.values():
                        if any(var in skill_lower for var in var_list):
                            variations.extend(var_list)
                    # Add the original skill
                    variations.append(skill_lower)
                    # Remove duplicates
                    variations = list(set(variations))
                    
                    for var in variations:
                        skill_texts.append(var)
                        skill_map[var] = skill
                
                # Convert to embeddings
                skill_embeddings = self.embedding_model.encode(skill_texts)
                resume_embeddings = self.embedding_model.encode(resume_sentences)
                
                # Calculate similarities
                similarities = resume_embeddings @ skill_embeddings.T
                max_similarities = np.max(similarities, axis=0)
                
                # Use a moderate threshold for semantic matching
                threshold = 0.6  # Higher threshold for more precise matching
                
                matched_variations = set()
                for idx, skill_text in enumerate(skill_texts):
                    if max_similarities[idx] > threshold:
                        original_skill = skill_map[skill_text]
                        if original_skill not in matched_skills:
                            matched_skills.append(original_skill)
                            matched_variations.add(skill_text)
            
            # Get missing skills
            missing_skills = [skill for skill in job_requirements['skills'] if skill not in matched_skills]
            
            return {
                'matched_skills': matched_skills,
                'missing_skills': missing_skills
            }
        except Exception as e:
            st.error(f"Error in skill matching: {str(e)}")
            return {
                'matched_skills': [],
                'missing_skills': []
            }
    def generate_cover_letter(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> str:
        """Generate a personalized cover letter based on the resume and job requirements."""
        prompt = f"""Write a professional cover letter for a job application following this specific format and guidelines.
        
        REQUIREMENTS:
        1. STRUCTURE:
           [Your Full Name]
           [Your Address]
           [City, State ZIP]
           [Email]
           [Phone]
           
           [Date]
           
           [Hiring Manager's Name/Title]
           [Company Name]
           [Company Address]
           [City, State ZIP]
           
           Dear [Hiring Manager's Name/Title],
           
           [Body Paragraphs]
           
           Sincerely,
           [Your Name]
        
        2. CONTENT GUIDELINES:
           - Opening: Express enthusiasm for the role and company
           - Body Paragraph 1: Match your skills to job requirements
           - Body Paragraph 2: Specific achievements that demonstrate value
           - Body Paragraph 3: Company knowledge and cultural fit
           - Closing: Clear call to action
        
        CONTEXT:
        Role: {job_requirements.get('title', '[Position]')}
        Company: {job_requirements.get('company', '[Company]')}
        Required Skills: {', '.join(job_requirements.get('skills', [])[:5])}
        Experience Needed: {job_requirements.get('experience', 'Not specified')}
        Your Matched Skills: {', '.join(skill_matches['matched_skills'])}
        
        Resume Context:
        {resume_text}

        Generate a complete cover letter following the exact format above. Keep it concise and focused on key achievements and relevant skills.
        """
        
        response = self.llm.invoke(prompt)
        return str(response.content)
    
    def tailor_resume(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Generate a tailored resume using the LLM and provide improvement analysis."""
        resume_prompt = f"""You are an expert ATS optimization specialist. Rewrite the following resume to maximize its ATS score while maintaining readability.
        The goal is to significantly improve the resume's ATS score by incorporating job-specific keywords and requirements.
        Keep in mind that dont add any skills that are not explicitly mentioned in the job requirements.
        keep in mind that dont add any of those things in the resume which are not already present in the resume.

        Job Requirements to Target:
        1. Required Skills: {job_requirements.get('skills', [])}
        2. Experience Needed: {job_requirements.get('experience', 'Not specified')}
        3. Responsibilities: {job_requirements.get('responsibilities', [])}

        Current Status:
        - Matched Skills: {skill_matches['matched_skills']}
        - Missing Skills: {skill_matches['missing_skills']}

        Optimization Requirements:
        1. Keyword Integration:
           - Add ALL missing required skills with relevant context
           - Place important keywords in prominent positions
           - Use exact phrases from job requirements
           - Maintain optimal keyword density (5-8%)
        
        2. Format Optimization:
           - Use clear section headers: Summary, Experience, Skills, Education
           - Start bullets with strong action verbs
           - Ensure consistent formatting
           - Use standard bullet points
        
        3. Content Enhancement:
           - Add quantifiable metrics to achievements
           - Highlight experience matching job requirements
           - Emphasize transferable skills
           - Use industry-standard terminology
        
        4. ATS Guidelines:
           - Use full terms before abbreviations
           - Avoid tables, columns, and graphics
           - Use standard job titles
           - Place keywords near the start of bullet points

        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Return ONLY the optimized resume text. Ensure EVERY required skill and responsibility is addressed.
        """
        
        tailored_resume = str(self.llm.invoke(resume_prompt).content)
        
        # Then, get the analysis separately
        analysis_prompt = f"""Analyze how the resume matches the job requirements and provide a detailed improvement analysis.
        
        IMPORTANT RULES:
        1. ONLY mention skills that are EXPLICITLY stated in the resume
        2. DO NOT make assumptions about skills not directly mentioned
        3. DO NOT infer skills from project descriptions unless explicitly stated
        4. If a skill is missing, list it in missing skills, do not try to find similar alternatives
        5. For matched skills, quote the exact text from resume that demonstrates the skill
        
        Focus on these aspects:
        1. Skills alignment - EXACT matches only
        2. Experience relevance - DIRECT matches only
        3. Achievement emphasis - ACTUAL achievements mentioned
        4. Missing keywords - List ALL required skills not found in resume
        
        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Currently Matched Skills (verified): {skill_matches['matched_skills']}
        Currently Missing Skills (verified): {skill_matches['missing_skills']}
        
        Return a JSON object with this exact structure:
        {{
            "improvements": [
                "specific improvements needed based on ACTUAL gaps"
            ],
            "skills_analysis": {{
                "matched": [
                    "ONLY skills explicitly found in resume with exact quotes"
                ],
                "missing": [
                    "ONLY skills from job requirements that are completely absent from resume"
                ]
            }},
            "achievement_emphasis": [
                "ONLY quantifiable achievements actually present in resume"
            ],
            "keyword_optimization": [
                "ONLY keywords from job requirements that should be added"
            ]
        }}
        
        Remember: Do not infer, assume, or suggest skills that are not explicitly stated in the resume.
        """
        
        analysis_response = self.llm.invoke(analysis_prompt)
        try:
            content = str(analysis_response.content).strip()
            # More robust JSON extraction
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            content = content.strip()
            
            try:
                analysis_result = json.loads(content)
                analysis_result["tailored_resume"] = tailored_resume
                return analysis_result
            except json.JSONDecodeError as json_err:
                return {
                    "improvements": ["Error analyzing improvements"],
                    "skills_analysis": {
                        "matched": ["Error analyzing matched skills"],
                        "missing": ["Error analyzing missing skills"]
                    },
                    "achievement_emphasis": ["Error analyzing achievements"],
                    "keyword_optimization": ["Error analyzing keywords"],
                    "tailored_resume": tailored_resume
                }
        except Exception as e:
            return {
                "improvements": ["Error analyzing improvements"],
                "skills_analysis": {
                    "matched": ["Error analyzing matched skills"],
                    "missing": ["Error analyzing missing skills"]
                },
                "achievement_emphasis": ["Error analyzing achievements"],
                "keyword_optimization": ["Error analyzing keywords"],
                "tailored_resume": tailored_resume
            }
    
    def generate_docx(self, tailored_content: str) -> Document:
        """Convert the tailored content into a DOCX file."""
        doc = Document()
        # Only include the actual resume content, no analysis
        content = str(tailored_content).strip()
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        return doc
    
    def generate_cold_email(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> str:
        """Generate a personalized cold email based on the resume and job requirements."""
        prompt = f"""Write a concise, professional cold email for a job application using the following details.
        Format exactly as shown, with clear sections and no commentary.
        
        REQUIREMENTS:
        1. FORMAT:
           - Subject line: Specific role and 1-2 key qualifications
           - Greeting: Professional and personalized
           - Body: 3-4 short paragraphs
           - Closing: Professional with clear call to action
        
        2. CONTENT GUIDELINES:
           - First paragraph: Introduction and mention 2-3 most relevant matching skills
           - Second paragraph: ONE specific, quantified achievement
           - Final paragraph: Brief call to action
           - Maximum 200 words
        
        CONTEXT:
        Position: {job_requirements.get('title', '[Position]')}
        Company: {job_requirements.get('company', '[Company]')}
        Key Requirements: {', '.join(job_requirements.get('skills', [])[:5])}
        Matched Skills: {', '.join(skill_matches['matched_skills'][:5])}
        
        Resume Details:
        {resume_text}
        
        EXAMPLE FORMAT:
        Subject: [Role] Application - [Key Qualification]
        
        Dear [Name],
        
        [Email Body]
        
        Best regards,
        [Full Name]
        [Contact Info]
        """
        
        response = self.llm.invoke(prompt)
        return str(response.content)

    def calculate_ats_score(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Calculate a comprehensive ATS score for the resume based on multiple factors."""
        prompt = f"""You are an ATS (Applicant Tracking System) expert. Analyze the resume against the job requirements and calculate scores.

        Rules for scoring:
        1. All scores must be integers (whole numbers)
        2. Each section score must not exceed its maximum value
        3. Total score must be the sum of all section scores
        4. Compare directly against job requirements
        5. Higher scores for exact keyword matches from requirements

        Scoring Criteria:
        1. Keyword Match (30 points max):
           - Award points for exact matches with job requirements
           - Check keyword frequency and placement
           - Required skills present: {job_requirements.get('skills', [])}
           - Current matched skills: {skill_matches['matched_skills']}

        2. Experience Alignment (25 points max):
           - Compare against required experience: {job_requirements.get('experience', 'Not specified')}
           - Check for relevant role titles
           - Evaluate described responsibilities against: {job_requirements.get('responsibilities', [])}

        3. Skills Match (25 points max):
           - Technical skills alignment
           - Soft skills presence
           - Skills context and application

        4. Education Relevance (10 points max):
           - Required education level match
           - Field of study relevance
           - Certifications value

        5. Format & Organization (10 points max):
           - Standard section headers
           - Bullet point structure
           - Content readability

        Resume to analyze:
        {resume_text}

        Job Requirements:
        {json.dumps(job_requirements, indent=2)}

        Return a JSON object with this exact structure:
        {{
            "total_score": <integer 0-100>,
            "section_scores": {{
                "keyword_match": {{
                    "score": <integer 0-30>,
                    "max": 30,
                    "details": ["<specific keywords found>", "<specific keywords missing>"]
                }},
                "experience": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<specific experience matches>", "<experience gaps>"]
                }},
                "skills": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<matched skills details>", "<missing skills impact>"]
                }},
                "education": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<education alignment details>"]
                }},
                "format": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<format strengths>", "<format improvements needed>"]
                }}
            }},
            "improvement_suggestions": [
                "<actionable suggestion 1>",
                "<actionable suggestion 2>",
                "<actionable suggestion 3>"
            ],
            "keyword_density": {{
                "<actual keyword from job requirements>": <integer frequency>
            }}
        }}
        """
        
        try:
            response = self.llm.invoke(prompt)
            content = str(response.content).strip()
            # Extract JSON content
            content = self._extract_json_content(content)
            
            # Parse and validate
            result = self._validate_ats_score(json.loads(content))
            return result
            
        except Exception as e:
            st.error(f"Error in ATS scoring: {str(e)}")
            return self._get_default_ats_score()
    
    def _extract_json_content(self, content: str) -> str:
        """Extract JSON content from LLM response."""
        if "```json" in content:
            content = content.split("```json")[1].split("```")[0]
        elif "```" in content:
            content = content.split("```")[1].split("```")[0]
        return content.strip()
    
    def _validate_ats_score(self, score_data: Dict) -> Dict:
        """Validate and fix ATS score data."""
        # Validate scores are integers and within range
        score_data["total_score"] = int(score_data["total_score"])
        
        # Ensure section scores are valid
        max_scores = {
            'keyword_match': 30,
            'experience': 25,
            'skills': 25,
            'education': 10,
            'format': 10
        }
        
        for section, max_score in max_scores.items():
            if section in score_data["section_scores"]:
                data = score_data["section_scores"][section]
                data["score"] = min(int(data["score"]), max_score)
                data["max"] = max_score
        
        # Recalculate total score
        total = sum(data["score"] for data in score_data["section_scores"].values())
        score_data["total_score"] = total
        
        return score_data
    
    def _get_default_ats_score(self) -> Dict:
        """Return default ATS score structure for error cases."""
        return {
            "total_score": 0,
            "section_scores": {
                "keyword_match": {"score": 0, "max": 30, "details": ["Error analyzing keywords"]},
                "experience": {"score": 0, "max": 25, "details": ["Error analyzing experience"]},
                "skills": {"score": 0, "max": 25, "details": ["Error analyzing skills"]},
                "education": {"score": 0, "max": 10, "details": ["Error analyzing education"]},
                "format": {"score": 0, "max": 10, "details": ["Error analyzing format"]}
            },
            "improvement_suggestions": ["Unable to generate suggestions due to an error"],
            "keyword_density": {}
        }

class AppNavigation:
    """Handle navigation and workflow selection in the app."""
    
    @staticmethod
    def show_landing_page():
        """Display the landing page with workflow options."""
        st.markdown("""
            <h1 style='text-align: center; color: #1E88E5; padding: 1rem;'>
                🎯 AI-Powered Resume Builder & Tailor
            </h1>
            <p style='text-align: center; color: #666; margin-bottom: 2rem;'>
                Create or optimize your resume with AI assistance
            </p>
        """, unsafe_allow_html=True)
        
        # Create a container for the workflow selection
        with st.container():
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("""
                    <div style='text-align: center; padding: 20px; border: 2px solid #1E88E5; border-radius: 10px;'>
                        <h2>🆕 Create New Resume</h2>
                        <p>Build a professional resume from scratch with AI guidance</p>
                    </div>
                """, unsafe_allow_html=True)
                if st.button("Start New Resume", key="new_resume", use_container_width=True):
                    st.session_state.workflow = "create_new"
            
            with col2:
                st.markdown("""
                    <div style='text-align: center; padding: 20px; border: 2px solid #1E88E5; border-radius: 10px;'>
                        <h2>✨ Tailor Existing Resume</h2>
                        <p>Optimize your existing resume for specific job positions</p>
                    </div>
                """, unsafe_allow_html=True)
                if st.button("Tailor Existing Resume", key="tailor_existing", use_container_width=True):
                    st.session_state.workflow = "tailor_existing"

def render_landing_page():
    """Render the landing page with workflow options."""
    st.markdown("""
        <h1 style='text-align: center; color: #1E88E5; padding: 1rem;'>
            🎯 AI-Powered Resume Builder & Tailor
        </h1>
            Create or optimize your resume with AI assistance
        </p>
    """, unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("""
            <div style='text-align: center; padding: 20px; border: 2px solid #1E88E5; border-radius: 10px;'>
                <h2>🆕 Create New Resume</h2>
                <p>Build a professional resume from scratch with AI guidance</p>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Start New Resume", key="new_resume", use_container_width=True):
            st.session_state.workflow = "create_new"
            st.rerun()
    
    with col2:
        st.markdown("""
            <div style='text-align: center; padding: 20px; border: 2px solid #1E88E5; border-radius: 10px;'>
                <h2>✨ Tailor Existing Resume</h2>
                <p>Optimize your existing resume for specific job positions</p>
            </div>
        """, unsafe_allow_html=True)
        if st.button("Tailor Existing Resume", key="tailor_existing", use_container_width=True):
            st.session_state.workflow = "tailor_existing"
            st.rerun()

def handle_create_new_workflow():
    """Handle the create new resume workflow."""
    # Initialize form data if needed
    if 'form_step' not in st.session_state:
        st.session_state.form_step = 1
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {
            'personal_info': {},
            'education': [],
            'experience': [],
            'skills': {'technical': [], 'soft': []},
            'certifications': []
        }
    
    form = ResumeForm()
    generator = ResumeGenerator()
    
    # Show progress bar
    form.render_progress_bar()
    
    # Handle form steps
    if st.session_state.form_step == 1:
        form.personal_info_form()
    elif st.session_state.form_step == 2:
        form.education_form()
    elif st.session_state.form_step == 3:
        form.experience_form()
    elif st.session_state.form_step == 4:
        form.skills_form()
    elif st.session_state.form_step == 5:
        if form.review_form():
            try:
                # Format data for template
                resume_data = generator.format_resume_data(st.session_state.form_data)
                
                # Generate resume versions
                html_content = generator.render_html(resume_data)
                pdf_content = generator.generate_pdf(html_content)
                
                # Show preview and download options
                preview_tabs = st.tabs(["📄 Resume Preview", "💾 Download Options"])
                
                with preview_tabs[0]:
                    st.components.v1.html(html_content, height=800, scrolling=True)
                
                with preview_tabs[1]:
                    st.download_button(
                        label="⬇️ Download PDF Resume",
                        data=pdf_content,
                        file_name="generated_resume.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    st.download_button(
                        label="⬇️ Download HTML Resume",
                        data=html_content,
                        file_name="generated_resume.html",
                        mime="text/html",
                        use_container_width=True
                    )
                
                # Show success message
                WorkflowManager.show_success_message("Resume generated successfully!")
                
            except Exception as e:
                st.error(f"Error generating resume: {str(e)}")

def initialize_session_state():
    """Initialize session state variables."""
    if 'workflow' not in st.session_state:
        st.session_state.workflow = "landing"
    if 'tailor_results' not in st.session_state:
        st.session_state.tailor_results = None

def store_tailor_results(results: dict):
    """Store all tailoring results in session state."""
    st.session_state.tailor_results = {
        'analysis_result': results['analysis_result'],
        'initial_ats_score': results['initial_ats_score'],
        'final_ats_score': results['final_ats_score'],
        'cold_email': results['cold_email'],
        'resume_text': results['resume_text'],
        'job_requirements': results['job_requirements'],
        'skill_matches': results['skill_matches']
    }

def get_stored_results():
    """Retrieve stored results from session state."""
    return st.session_state.tailor_results

def handle_tailor_workflow():
    """Handle the resume tailoring workflow."""
    try:
        tailor = ResumeTailor()
        
        left_col, right_col = st.columns([1, 1.5], gap="large")
        
        with left_col:
            resume_text = WorkflowManager.handle_file_upload(['pdf', 'docx'])
            job_text, is_url = WorkflowManager.handle_job_input()
            
            if resume_text and job_text:
                process_button = st.button("🎯 Tailor Resume", use_container_width=True)
            else:
                st.info("Please upload your resume and provide job details to proceed.")
                process_button = False
        
        with right_col:
            if resume_text and job_text and process_button:
                with st.spinner("Analyzing and optimizing your resume..."):
                    # Parse job description
                    job_requirements = tailor.parse_job_description(job_text, is_url)
                    
                    # Match skills
                    skill_matches = tailor.match_skills(resume_text, job_requirements)
                    
                    # Calculate initial ATS score
                    initial_ats_score = tailor.calculate_ats_score(resume_text, job_requirements, skill_matches)
                    
                    # Generate tailored resume with analysis
                    analysis_result = tailor.tailor_resume(resume_text, job_requirements, skill_matches)
                    
                    # Add job requirements and skill matches to analysis result
                    analysis_result['job_requirements'] = job_requirements
                    analysis_result['skill_matches'] = skill_matches
                    
                    # Calculate final ATS score
                    final_ats_score = tailor.calculate_ats_score(
                        analysis_result['tailored_resume'], 
                        job_requirements, 
                        skill_matches
                    )
                    
                    # Generate cold email
                    cold_email = tailor.generate_cold_email(resume_text, job_requirements, skill_matches)
                    
                    # Store results in session state
                    store_tailor_results({
                        'analysis_result': analysis_result,
                        'initial_ats_score': initial_ats_score,
                        'final_ats_score': final_ats_score,
                        'cold_email': cold_email,
                        'resume_text': resume_text,
                        'job_requirements': job_requirements,
                        'skill_matches': skill_matches
                    })
                    
                    # Show results
                    show_results_tabs(
                        analysis_result,
                        initial_ats_score,
                        final_ats_score,
                        cold_email,
                        resume_text,
                        tailor
                    )
    
    except ValueError as e:
        st.error(f"Configuration Error: {str(e)}")
        st.info("Please ensure you have set up the GROQ_API_KEY in your .env file")

def show_results_tabs(analysis_result, initial_ats_score, final_ats_score, cold_email, resume_text, tailor):
    """Show the results in organized tabs."""
    try:
        # Get job requirements and skill matches from the analysis result
        job_requirements = analysis_result.get('job_requirements', {})
        skill_matches = analysis_result.get('skill_matches', {'matched_skills': [], 'missing_skills': []})
        
        # Create tabs
        tabs = st.tabs(["💡 Analysis", "📊 ATS Score", "📧 Cold Email", "📝 Cover Letter", "📄 Resume Versions"])
        
        # Analysis Tab
        with tabs[0]:
            show_analysis_tab(analysis_result)
        
        # ATS Score Tab
        with tabs[1]:
            show_ats_score_tab(initial_ats_score, final_ats_score)
        
        # Cold Email Tab
        with tabs[2]:
            st.text_area("", cold_email, height=400)
        
        # Cover Letter Tab
        with tabs[3]:
            with st.spinner("Generating cover letter..."):
                cover_letter = tailor.generate_cover_letter(resume_text, job_requirements, skill_matches)
                st.text_area("", cover_letter, height=600)
                if cover_letter:
                    st.download_button(
                        label="⬇️ Download Cover Letter",
                        data=cover_letter,
                        file_name="cover_letter.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
        
        # Resume Versions Tab
        with tabs[4]:
            show_resume_versions_tab(resume_text, analysis_result, tailor)
            
    except Exception as e:
        st.error(f"Error displaying results: {str(e)}")

def show_analysis_tab(analysis_result):
    """Display the analysis tab content."""
    # Skills Analysis
    st.subheader("🎯 Skills Analysis")
    skill_cols = st.columns(2)
    with skill_cols[0]:
        st.write("✅ Matched Skills")
        for idx, skill_analysis in enumerate(analysis_result['skills_analysis']['matched'], 1):
            st.write(f"{idx}. {skill_analysis}")
    with skill_cols[1]:
        st.write("📝 Missing Skills & Suggestions")
        for idx, skill_suggestion in enumerate(analysis_result['skills_analysis']['missing'], 1):
            st.write(f"{idx}. {skill_suggestion}")
    
    # Improvements
    st.subheader("🔄 Improvements Made")
    for idx, improvement in enumerate(analysis_result['improvements'], 1):
        st.write(f"{idx}. {improvement}")
    
    # Achievements
    st.subheader("📊 Quantifiable Achievements")
    for idx, achievement in enumerate(analysis_result['achievement_emphasis'], 1):
        st.write(f"{idx}. {achievement}")
    
    # Keywords
    st.subheader("🔍 ATS Keywords")
    for idx, keyword in enumerate(analysis_result['keyword_optimization'], 1):
        st.write(f"{idx}. {keyword}")

def validate_ats_score(score_data: Dict) -> Dict:
    """Validate and fix ATS score data."""
    try:
        # Ensure all sections exist
        sections = ['keyword_match', 'experience', 'skills', 'education', 'format']
        max_scores = {'keyword_match': 30, 'experience': 25, 'skills': 25, 'education': 10, 'format': 10}
        
        for section in sections:
            if section not in score_data['section_scores']:
                score_data['section_scores'][section] = {
                    'score': 0,
                    'max': max_scores[section],
                    'details': ['Section not evaluated']
                }
            else:
                # Ensure score doesn't exceed maximum
                section_data = score_data['section_scores'][section]
                section_data['max'] = max_scores[section]
                section_data['score'] = min(int(section_data['score']), max_scores[section])
        
        # Recalculate total score
        total = sum(data['score'] for data in score_data['section_scores'].values())
        score_data['total_score'] = total
        
        return score_data
        
    except Exception as e:
        st.error(f"Error validating ATS score: {str(e)}")
        return score_data

def show_ats_score_tab(initial_ats_score, final_ats_score):
    """Display the ATS score tab content with high contrast."""
    try:
        # Validate scores
        initial_ats_score = validate_ats_score(initial_ats_score)
        final_ats_score = validate_ats_score(final_ats_score)
        
        # Container for better spacing and contrast
        with st.container():
            st.markdown("""
                <div style='background-color: #ffffff; padding: 20px; border-radius: 10px; border: 2px solid #000000;'>
                    <h2 style='text-align: center; color: #000000; margin-bottom: 30px;'>ATS Score Analysis</h2>
                </div>
            """, unsafe_allow_html=True)
            
            # Display scores side by side
            col1, col2 = st.columns(2)
            
            # Before Score
            with col1:
                initial_total = initial_ats_score["total_score"]
                score_text = "Needs Improvement"
                text_color = "#800000"  # Dark red for contrast
                if initial_total >= 70:
                    score_text = "Excellent"
                    text_color = "#006400"  # Dark green
                elif initial_total >= 50:
                    score_text = "Good"
                    text_color = "#000000"  # Black
                
                st.markdown(
                    f"""
                    <div style='text-align: center; padding: 20px; border: 2px solid #000000; border-radius: 10px; background-color: #ffffff;'>
                        <h3 style='color: #000000; margin-bottom: 10px; font-weight: bold;'>Original Score</h3>
                        <div style='font-size: 48px; margin: 20px 0; font-weight: bold; color: {text_color};'>{initial_total}</div>
                        <p style='color: {text_color}; font-size: 18px; font-weight: bold;'>{score_text}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
            # After Score
            with col2:
                final_total = final_ats_score["total_score"]
                score_text = "Needs Improvement"
                text_color = "#800000"  # Dark red for contrast
                if final_total >= 70:
                    score_text = "Excellent"
                    text_color = "#006400"  # Dark green
                elif final_total >= 50:
                    score_text = "Good"
                    text_color = "#000000"  # Black
                
                st.markdown(
                    f"""
                    <div style='text-align: center; padding: 20px; border: 2px solid #000000; border-radius: 10px; background-color: #ffffff;'>
                        <h3 style='color: #000000; margin-bottom: 10px; font-weight: bold;'>Optimized Score</h3>
                        <div style='font-size: 48px; margin: 20px 0; font-weight: bold; color: {text_color};'>{final_total}</div>
                        <p style='color: {text_color}; font-size: 18px; font-weight: bold;'>{score_text}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
            # Improvement percentage with better contrast
            if initial_total > 0:
                improvement = ((final_total - initial_total) / initial_total) * 100
                improvement_color = "#006400" if improvement > 0 else "#800000"
                st.markdown(
                    f"""
                    <div style='text-align: center; margin: 30px 0; padding: 15px; background-color: #ffffff; border: 2px solid #000000; border-radius: 10px;'>
                        <h3 style='color: {improvement_color}; font-size: 24px; font-weight: bold;'>
                            Overall Improvement: {'+' if improvement > 0 else ''}{improvement:.1f}%
                        </h3>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
            
            # Section Scores with better contrast
            st.markdown("""
                <div style='margin: 30px 0;'>
                    <h3 style='color: #000000; font-weight: bold; margin-bottom: 20px;'>Detailed Section Analysis</h3>
                </div>
            """, unsafe_allow_html=True)
            
            # Display each section score
            for section, data in final_ats_score["section_scores"].items():
                section_name = section.replace('_', ' ').title()
                initial_score = initial_ats_score["section_scores"][section]["score"]
                final_score = data["score"]
                max_score = data["max"]
                
                st.markdown(
                    f"""
                    <div style='background-color: #ffffff; padding: 15px; border: 1px solid #000000; border-radius: 5px; margin-bottom: 15px;'>
                        <h4 style='color: #000000; margin-bottom: 15px; font-weight: bold;'>{section_name}</h4>
                    """,
                    unsafe_allow_html=True
                )
                
                cols = st.columns([4, 4, 2])
                
                # Before score
                with cols[0]:
                    initial_percent = (initial_score / max_score) * 100
                    st.markdown(
                        f"""
                        <div style='padding: 10px; border: 1px solid #000000; border-radius: 5px;'>
                            <p style='color: #000000; margin-bottom: 5px; font-weight: bold;'>Before: {initial_score}/{max_score}</p>
                            <div style='background-color: #e0e0e0; height: 10px; border-radius: 5px;'>
                                <div style='background-color: #000000; width: {initial_percent}%; height: 100%; border-radius: 5px;'></div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                
                # After score
                with cols[1]:
                    final_percent = (final_score / max_score) * 100
                    st.markdown(
                        f"""
                        <div style='padding: 10px; border: 1px solid #000000; border-radius: 5px;'>
                            <p style='color: #000000; margin-bottom: 5px; font-weight: bold;'>After: {final_score}/{max_score}</p>
                            <div style='background-color: #e0e0e0; height: 10px; border-radius: 5px;'>
                                <div style='background-color: #000000; width: {final_percent}%; height: 100%; border-radius: 5px;'></div>
                            </div>
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                
                # Details button
                with cols[2]:
                    if st.button("👉 Details", key=f"details_{section}"):
                        st.markdown("""
                            <div style='background-color: #ffffff; padding: 15px; border: 1px solid #000000; border-radius: 5px; margin-top: 10px;'>
                        """, unsafe_allow_html=True)
                        st.markdown("**Before:**")
                        for detail in initial_ats_score["section_scores"][section]["details"]:
                            st.markdown(f"• {detail}")
                        st.markdown("**After:**")
                        for detail in data["details"]:
                            st.markdown(f"• {detail}")
                        st.markdown("</div>", unsafe_allow_html=True)
                
                st.markdown("</div>", unsafe_allow_html=True)
            
            # Keyword Density Analysis with better contrast
            st.markdown("""
                <div style='margin: 30px 0;'>
                    <h3 style='color: #000000; font-weight: bold;'>Keyword Density Comparison</h3>
                </div>
            """, unsafe_allow_html=True)
            
            density_cols = st.columns(2)
            
            with density_cols[0]:
                st.markdown("<p style='color: #000000; font-weight: bold;'>Before Optimization</p>", unsafe_allow_html=True)
                if initial_ats_score["keyword_density"]:
                    df = pd.DataFrame(
                        list(initial_ats_score["keyword_density"].items()),
                        columns=["Keyword", "Frequency"]
                    ).sort_values(by="Frequency", ascending=False)
                    st.dataframe(
                        df,
                        hide_index=True,
                        use_container_width=True
                    )
            
            with density_cols[1]:
                st.markdown("<p style='color: #000000; font-weight: bold;'>After Optimization</p>", unsafe_allow_html=True)
                if final_ats_score["keyword_density"]:
                    df = pd.DataFrame(
                        list(final_ats_score["keyword_density"].items()),
                        columns=["Keyword", "Frequency"]
                    ).sort_values(by="Frequency", ascending=False)
                    st.dataframe(
                        df,
                        hide_index=True,
                        use_container_width=True
                    )
            
            # Improvement Suggestions with better contrast
            st.markdown("""
                <div style='margin: 30px 0;'>
                    <h3 style='color: #000000; font-weight: bold;'>Improvement Suggestions</h3>
                </div>
            """, unsafe_allow_html=True)
            
            for idx, suggestion in enumerate(final_ats_score["improvement_suggestions"], 1):
                st.markdown(
                    f"""
                    <div style='background-color: #ffffff; padding: 15px; border: 1px solid #000000; border-radius: 5px; margin-bottom: 10px;'>
                        <p style='color: #000000; margin: 0; font-weight: 500;'>{idx}. {suggestion}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )
                
    except Exception as e:
        st.error(f"Error displaying ATS score: {str(e)}")
        st.markdown(
            """
            <div style='background-color: #ffffff; padding: 15px; border: 2px solid #800000; border-radius: 5px;'>
                <p style='color: #800000; margin: 0; font-weight: bold;'>Unable to display ATS score analysis. Please try again.</p>
            </div>
            """,
            unsafe_allow_html=True
        )

def show_resume_versions_tab(resume_text, analysis_result, tailor):
    """Display the resume versions tab content with professional template options."""
    # Store resume data in session state if not already stored
    if 'generated_resume_data' not in st.session_state:
        try:
            # Extract resume sections from optimized content
            optimized_content = analysis_result['tailored_resume'].split("Here's the optimized resume:")[-1].strip() if "Here's the optimized resume:" in analysis_result['tailored_resume'] else analysis_result['tailored_resume']
            sections = parse_resume_sections(optimized_content)
            
            # Store data in session state
            st.session_state.generated_resume_data = {
                'original_text': resume_text,
                'optimized_content': optimized_content,
                'sections': sections,
                'html_content': None,
                'pdf_content': None,
                'doc': None
            }
            
            # Initialize ResumeGenerator
            generator = ResumeGenerator()
            
            # Generate formatted resume using template
            resume_data = {
                'personal_info': sections['personal_info'],
                'education': sections['education'],
                'experience': sections['experience'],
                'skills': sections['skills'],
                'certifications': sections.get('certifications', [])
            }
            
            # Generate different formats and store them
            st.session_state.generated_resume_data['html_content'] = generator.render_html(resume_data)
            st.session_state.generated_resume_data['pdf_content'] = generator.generate_pdf(st.session_state.generated_resume_data['html_content'])
            st.session_state.generated_resume_data['doc'] = generator.generate_docx(resume_data)
        
        except KeyError as e:
            st.error(f"Missing key in resume data: {e}")
        except AttributeError as e:
            st.error(f"Unexpected data format: {e}")
        except Exception as e:
            st.error(f"An error occurred while generating resume versions: {e}")

            
    # Display the content
    version_cols = st.columns(2)
    
    with version_cols[0]:
        st.markdown("#### Original Resume")
        st.text_area("", st.session_state.generated_resume_data['original_text'], height=400, disabled=True)
    
    with version_cols[1]:
        st.markdown("#### Optimized Resume")
        st.text_area("", st.session_state.generated_resume_data['optimized_content'], height=400, disabled=True)
    
    # Add professional template generation
    st.markdown("---")
    st.markdown("### Download Professional Resume")
    
    try:
        # Show preview with improved styling
        st.markdown("#### Preview")
        st.markdown("""
            <div style='padding: 20px; background-color: #ffffff; border: 1px solid #ddd; border-radius: 5px;'>
        """, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Error displaying preview: {str(e)}")
        st.components.v1.html(
            f"""
            <div style="background-color: white; padding: 20px;">
                {st.session_state.generated_resume_data['html_content']}
            </div>
            """,
            height=800,
            scrolling=True
        )
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Download options
        st.markdown("#### Download Options")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                label="⬇️ Download PDF Version",
                data=st.session_state.generated_resume_data['pdf_content'],
                file_name="tailored_resume.pdf",
                mime="application/pdf",
                use_container_width=True
            )
        
        with col2:
            st.download_button(
                label="⬇️ Download HTML Version",
                data=st.session_state.generated_resume_data['html_content'],
                file_name="tailored_resume.html",
                mime="text/html",
                use_container_width=True
            )
        
        with col3:
            # Save DOCX temporarily and read it back
            temp_docx = "tailored_resume.docx"
            st.session_state.generated_resume_data['doc'].save(temp_docx)
            with open(temp_docx, "rb") as file:
                docx_data = file.read()
            os.remove(temp_docx)  # Clean up the temporary file
            
            st.download_button(
                label="⬇️ Download Word Version",
                data=docx_data,
                file_name="tailored_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
    except Exception as e:
        st.error(f"Error displaying resume: {str(e)}")
        st.info("Falling back to basic format...")
        # Fallback to basic DOCX
        doc = tailor.generate_docx(st.session_state.generated_resume_data['optimized_content'])
        temp_docx = "tailored_resume.docx"
        doc.save(temp_docx)
        with open(temp_docx, "rb") as file:
            docx_data = file.read()
        os.remove(temp_docx)
        
        st.download_button(
            label="⬇️ Download Basic Resume",
            data=docx_data,
            file_name="tailored_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

def parse_resume_sections(content: str) -> Dict:
    """Parse the optimized resume content into structured sections."""
    try:
        sections = {
            'personal_info': {
                'full_name': '',
                'email': '',
                'phone': '',
                'location': '',
                'linkedin': '',
                'summary': ''
            },
            'education': [],
            'experience': [],
            'skills': {
                'technical': [],
                'soft': []
            },
            'certifications': []
        }
        
        # Split content into sections
        lines = content.split('\n')
        current_section = None
        temp_experience = {}
        
        for idx, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Extract personal info from first few lines
            if idx < 5:
                if '@' in line:
                    sections['personal_info']['email'] = line
                elif any(x.isdigit() for x in line):
                    sections['personal_info']['phone'] = line
                elif 'linkedin.com' in line.lower():
                    sections['personal_info']['linkedin'] = line
                else:
                    if not sections['personal_info']['full_name']:
                        sections['personal_info']['full_name'] = line
                    elif not sections['personal_info']['location']:
                        sections['personal_info']['location'] = line
            
            # Detect sections
            lower_line = line.lower()
            if 'summary' in lower_line or 'profile' in lower_line or 'objective' in lower_line:
                current_section = 'summary'
                continue
            elif 'education' in lower_line and len(line) < 20:
                current_section = 'education'
                sections['education'] = []
                continue
            elif 'experience' in lower_line and len(line) < 20:
                current_section = 'experience'
                sections['experience'] = []
                temp_experience = {}
                continue
            elif 'skills' in lower_line and len(line) < 20:
                current_section = 'skills'
                continue
            elif 'certification' in lower_line and len(line) < 20:
                current_section = 'certifications'
                continue
            
            # Process section content
            if current_section == 'summary':
                if not any(keyword in lower_line for keyword in ['education', 'experience', 'skills', 'certification']):
                    sections['personal_info']['summary'] += f"{line} "
            
            elif current_section == 'education':
                if line:
                    edu_entry = {
                        'degree': line,
                        'institution': '',
                        'start_year': '',
                        'end_year': ''
                    }
                    # Try to extract years
                    years = [word for word in line.split() if word.isdigit() and len(word) == 4]
                    if len(years) >= 2:
                        edu_entry['start_year'] = years[0]
                        edu_entry['end_year'] = years[1]
                    # Extract institution
                    if ',' in line:
                        parts = line.split(',')
                        edu_entry['institution'] = parts[0].strip()
                        edu_entry['degree'] = parts[1].strip()
                    sections['education'].append(edu_entry)
            
            elif current_section == 'experience':
                if line.strip():
                    if line.startswith('•') or line.startswith('-'):
                        if 'responsibilities' not in temp_experience:
                            temp_experience['responsibilities'] = []
                        temp_experience['responsibilities'].append(line.lstrip('•- ').strip())
                    else:
                        if temp_experience and 'title' in temp_experience:
                            sections['experience'].append(temp_experience)
                            temp_experience = {}
                        
                        # New experience entry
                        title_parts = line.split(',')
                        temp_experience['title'] = title_parts[0].strip()
                        if len(title_parts) > 1:
                            temp_experience['company'] = title_parts[1].strip()
                        
                        # Try to extract dates
                        date_parts = [part for part in line.split() if '-' in part or 'present' in part.lower()]
                        if date_parts:
                            dates = date_parts[0].split('-')
                            temp_experience['start_date'] = dates[0].strip()
                            temp_experience['end_date'] = dates[1].strip() if len(dates) > 1 else 'Present'
            
            elif current_section == 'skills':
                if 'technical' in lower_line:
                    sections['skills']['technical'] = [s.strip() for s in line.split(':')[1].split(',') if s.strip()]
                elif 'soft' in lower_line:
                    sections['skills']['soft'] = [s.strip() for s in line.split(':')[1].split(',') if s.strip()]
                elif line:
                    if not sections['skills']['technical']:
                        sections['skills']['technical'] = [s.strip() for s in line.split(',') if s.strip()]
            
            elif current_section == 'certifications':
                if line:
                    cert_parts = line.split(',')
                    cert_entry = {
                        'name': cert_parts[0].strip(),
                        'issuer': cert_parts[1].strip() if len(cert_parts) > 1 else '',
                        'year': cert_parts[2].strip() if len(cert_parts) > 2 else ''
                    }
                    sections['certifications'].append(cert_entry)
        
        # Add last experience entry if exists
        if temp_experience and 'title' in temp_experience:
            sections['experience'].append(temp_experience)
        
        return sections
    
    except Exception as e:
        st.error(f"Error parsing resume sections: {str(e)}")
        return {
            'personal_info': {
                'full_name': 'Name not found',
                'email': 'Email not found',
                'phone': 'Phone not found',
                'location': 'Location not found',
                'linkedin': '',
                'summary': 'Summary not found'
            },
            'education': [],
            'experience': [],
            'skills': {'technical': [], 'soft': []},
            'certifications': []
        }

def load_css():
    """Load custom CSS styles with high contrast black and white theme."""
    st.markdown("""
        <style>
        /* Global Styles */
        [data-testid="stAppViewContainer"] {
            background: #ffffff;
        }
        
        /* Text Colors */
        h1, h2, h3, p, span, div, label, text {
            color: #000000 !important;
            font-family: 'Arial', sans-serif;
        }
        
        /* Buttons */
        div.stButton > button {
            background: #000000;
            color: #ffffff;
            border: none;
            padding: 0.75rem 1.5rem;
            border-radius: 4px;
            font-weight: 500;
        }
        
        /* Input Fields */
        div.stTextInput > div > div > input,
        div.stTextArea > div > textarea {
            color: #000000;
            background-color: #ffffff;
            border: 1px solid #000000;
        }
        
        /* Progress Bars */
        div.stProgress > div > div {
            background-color: #000000;
        }
        
        /* Cards and Sections */
        .css-1y4p8pa {
            background: #ffffff;
            border: 1px solid #000000;
            padding: 1rem;
        }
        
        /* Tabs */
        .stTabs [data-baseweb="tab"] {
            color: #000000;
        }
        
        .stTabs [aria-selected="true"] {
            background-color: #000000 !important;
            color: #ffffff !important;
        }
        
        /* Info Messages */
        .stInfo {
            color: #000000;
            background-color: #ffffff;
            border: 1px solid #000000;
        }
        
        /* Error Messages */
        .stError {
            color: #800000;
            background-color: #ffffff;
            border: 1px solid #800000;
        }
        
        /* Success Messages */
        .stSuccess {
            color: #004d00;
            background-color: #ffffff;
            border: 1px solid #004d00;
        }
        
        /* Tables and Dataframes */
        .dataframe {
            color: #000000;
        }
        
        /* Text Areas */
        .stTextArea textarea {
            color: #000000;
            background-color: #ffffff;
            border: 1px solid #000000;
        }
        
        /* Links */
        a {
            color: #000000;
            text-decoration: underline;
        }
        
        /* Radio Buttons */
        .stRadio label {
            color: #000000;
        }
        
        /* Expandable Sections */
        .streamlit-expanderHeader {
            color: #000000;
            background-color: #ffffff;
        }
        
        /* Preview iframe styling */
        iframe {
            background-color: white !important;
            padding: 20px !important;
        }
        
        /* Set dark text color for all text elements */
        .stMarkdown, .stText {
            color: #000000 !important;
        }
        </style>
    """, unsafe_allow_html=True)

def main():
    st.set_page_config(layout="wide", page_title="AI Resume Builder & Tailor")
    
    # Initialize session state
    initialize_session_state()
    
    # Handle routing
    if st.session_state.workflow == "landing":
        render_landing_page()
    else:
        # Show navigation
        col1, col2 = st.columns([1, 4])
        with col1:
            if st.button("← Back to Home"):
                # Reset all session state
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.session_state.workflow = "landing"
                st.rerun()
        
        with col2:
            if st.session_state.workflow == "create_new":
                st.markdown(f"**Current:** Create New Resume > Step {st.session_state.get('form_step', 1)}")
            elif st.session_state.workflow == "tailor_existing":
                st.markdown("**Current:** Tailor Existing Resume")
        
        # Handle workflows
        if st.session_state.workflow == "create_new":
            handle_create_new_workflow()
        elif st.session_state.workflow == "tailor_existing":
            handle_tailor_workflow()

if __name__ == "__main__":
    main()








